"""
Générateur de documents
Création de PDF, DOCX et autres formats avec Ollama
"""

import os
from datetime import datetime
from typing import Any, Dict, Optional, TYPE_CHECKING

import docx
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

# Import du LLM local (Ollama)
if TYPE_CHECKING:
    from models.local_llm import LocalLLM

try:
    from models.local_llm import LocalLLM
    OLLAMA_AVAILABLE = True
except ImportError:
    OLLAMA_AVAILABLE = False


class DocumentGenerator:
    """
    Générateur de documents dans différents formats avec génération IA
    """

    def __init__(self, llm: Optional[LocalLLM] = None):
        """
        Initialise le générateur de documents
        
        Args:
            llm: Instance de LocalLLM (Ollama) pour la génération de contenu
        """
        self.llm = llm if llm else (LocalLLM() if OLLAMA_AVAILABLE else None)
        self._check_dependencies()

    def _check_dependencies(self):
        """
        Vérifie la disponibilité des bibliothèques
        """
        self.reportlab_available = False
        self.python_docx_available = False

        try:
            self.reportlab_available = True
        except ImportError:
            pass

        try:
            self.python_docx_available = True
        except ImportError:
            pass

    async def generate_document(
        self, query: str, context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Génère un document basé sur la requête

        Args:
            query: Demande de génération
            context: Contexte et données

        Returns:
            Informations sur le document généré
        """
        # Analyse de la requête pour déterminer le format
        format_type = self._determine_format(query)

        if format_type == "pdf":
            return await self.generate_pdf(query, context)
        elif format_type == "docx":
            return await self.generate_docx(query, context)
        else:
            return await self.generate_text(query, context)

    def _determine_format(self, query: str) -> str:
        """
        Détermine le format de document à générer
        """
        query_lower = query.lower()

        if "pdf" in query_lower:
            return "pdf"
        elif "word" in query_lower or "docx" in query_lower:
            return "docx"
        else:
            return "text"

    async def generate_pdf(
        self, content: str, context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Génère un document PDF

        Args:
            content: Contenu à inclure
            context: Contexte et paramètres

        Returns:
            Résultat de la génération
        """
        if not self.reportlab_available:
            return {
                "error": "ReportLab non disponible. Installez avec: pip install reportlab",
                "success": False,
            }

        try:
            # Nom de fichier
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.pdf"
            filepath = os.path.join("outputs", filename)

            # Création du répertoire si nécessaire
            os.makedirs("outputs", exist_ok=True)

            # Configuration du document
            doc = SimpleDocTemplate(filepath, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # Titre
            title = context.get("title", "Document Généré par IA")
            story.append(Paragraph(title, styles["Title"]))
            story.append(Spacer(1, 12))

            # Métadonnées
            story.append(
                Paragraph(
                    f"Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}",
                    styles["Normal"],
                )
            )
            story.append(Spacer(1, 12))

            # Contenu principal
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    story.append(Paragraph(para.strip(), styles["Normal"]))
                    story.append(Spacer(1, 6))

            # Construction du PDF
            doc.build(story)

            return {
                "success": True,
                "file_path": filepath,
                "file_name": filename,
                "format": "pdf",
                "size": os.path.getsize(filepath),
            }

        except Exception as e:
            return {
                "error": f"Erreur lors de la génération PDF: {str(e)}",
                "success": False,
            }

    async def generate_docx(
        self, content: str, context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Génère un document Word (DOCX)

        Args:
            content: Contenu à inclure
            context: Contexte et paramètres

        Returns:
            Résultat de la génération
        """
        if not self.python_docx_available:
            return {
                "error": "python-docx non disponible. Installez avec: pip install python-docx",
                "success": False,
            }

        try:
            # Nom de fichier
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.docx"
            filepath = os.path.join("outputs", filename)

            # Création du répertoire si nécessaire
            os.makedirs("outputs", exist_ok=True)

            # Création du document
            doc = docx.Document()

            # Titre
            title = context.get("title", "Document Généré par IA")
            doc.add_heading(title, 0)

            # Métadonnées
            doc.add_paragraph(
                f"Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            )
            doc.add_paragraph("")  # Ligne vide

            # Contenu principal
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    # Détection des titres (lignes courtes en majuscules)
                    if len(para.strip()) < 100 and para.strip().isupper():
                        doc.add_heading(para.strip(), level=1)
                    else:
                        doc.add_paragraph(para.strip())

            # Sauvegarde
            doc.save(filepath)

            return {
                "success": True,
                "file_path": filepath,
                "file_name": filename,
                "format": "docx",
                "size": os.path.getsize(filepath),
            }

        except Exception as e:
            return {
                "error": f"Erreur lors de la génération DOCX: {str(e)}",
                "success": False,
            }

    async def generate_text(
        self, content: str, context: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Génère un fichier texte simple

        Args:
            content: Contenu à inclure
            context: Contexte et paramètres

        Returns:
            Résultat de la génération
        """
        try:
            # Nom de fichier
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.txt"
            filepath = os.path.join("outputs", filename)

            # Création du répertoire si nécessaire
            os.makedirs("outputs", exist_ok=True)

            # Contenu du fichier
            full_content = []

            # Titre
            title = context.get("title", "Document Généré par IA")
            full_content.append(title)
            full_content.append("=" * len(title))
            full_content.append("")

            # Métadonnées
            full_content.append(
                f"Généré le: {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
            )
            full_content.append("")

            # Contenu principal
            full_content.append(content)

            # Écriture du fichier
            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(full_content))

            return {
                "success": True,
                "file_path": filepath,
                "file_name": filename,
                "format": "text",
                "size": os.path.getsize(filepath),
            }

        except Exception as e:
            return {
                "error": f"Erreur lors de la génération du fichier texte: {str(e)}",
                "success": False,
            }

    async def generate_report(self, data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Génère un rapport structuré

        Args:
            data: Données pour le rapport
            template: Template à utiliser

        Returns:
            Rapport généré
        """
        try:
            # Formatage du contenu du rapport
            content_parts = []

            # En-tête
            content_parts.append("RAPPORT D'ANALYSE")
            content_parts.append("")

            # Résumé exécutif
            if "summary" in data:
                content_parts.append("RÉSUMÉ EXÉCUTIF")
                content_parts.append("-" * 16)
                content_parts.append(data["summary"])
                content_parts.append("")

            # Données principales
            if "main_data" in data:
                content_parts.append("DONNÉES PRINCIPALES")
                content_parts.append("-" * 18)
                for key, value in data["main_data"].items():
                    content_parts.append(f"{key}: {value}")
                content_parts.append("")

            # Analyses
            if "analysis" in data:
                content_parts.append("ANALYSE DÉTAILLÉE")
                content_parts.append("-" * 17)
                content_parts.append(data["analysis"])
                content_parts.append("")

            # Recommandations
            if "recommendations" in data:
                content_parts.append("RECOMMANDATIONS")
                content_parts.append("-" * 15)
                if isinstance(data["recommendations"], list):
                    for i, rec in enumerate(data["recommendations"], 1):
                        content_parts.append(f"{i}. {rec}")
                else:
                    content_parts.append(data["recommendations"])
                content_parts.append("")

            content = "\n".join(content_parts)

            # Génération du rapport
            context = {
                "title": "Rapport d'Analyse IA",
                "format": data.get("format", "pdf"),
            }

            return await self.generate_document(content, context)

        except Exception as e:
            return {
                "error": f"Erreur lors de la génération du rapport: {str(e)}",
                "success": False,
            }

    def get_dependencies_status(self) -> Dict[str, bool]:
        """
        Retourne le statut des dépendances
        """
        return {
            "reportlab": self.reportlab_available,
            "python_docx": self.python_docx_available,
            "text_generation": True,  # Toujours disponible
        }
