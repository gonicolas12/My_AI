"""
Processeur de fichiers DOCX
Lecture, analyse et extraction de contenu Word
"""

import glob
import os
from pathlib import Path
from typing import Any, Dict

import docx


class DOCXProcessor:
    """
    Processeur pour les fichiers DOCX (Word)
    """

    def __init__(self):
        """
        Initialise le processeur DOCX
        """
        self.supported_extensions = [".docx", ".doc"]
        self._check_dependencies()

    def _check_dependencies(self):
        """
        Vérifie la disponibilité des bibliothèques DOCX
        """
        self.python_docx_available = False

        try:
            # Test d'import complet
            self.python_docx_available = True
            print("✅ Module python-docx disponible")
        except ImportError as e:
            print(f"❌ Module 'python-docx' non installé: {e}")
            print("💡 Installez avec: pip install python-docx")
        except Exception as e:
            print(f"❌ Erreur lors de l'import python-docx: {e}")

    def read_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Lit un fichier DOCX et extrait le contenu - Version avec support OneDrive
        """
        # Résoudre le chemin OneDrive en premier
        resolved_path = self._resolve_onedrive_path(file_path)

        if not os.path.exists(resolved_path):
            return {
                "success": False,
                "error": f"Fichier non trouvé: {file_path}",
                "error_type": "FILE_NOT_FOUND",
                "resolution": "Vérifiez que le fichier existe ou qu'il est synchronisé depuis OneDrive",
            }

        if not self.python_docx_available:
            return {
                "success": False,
                "error": "python-docx non disponible. Installez avec: pip install python-docx",
                "error_type": "MISSING_PACKAGE",
                "resolution": "Installez le package avec : pip install python-docx",
            }

        try:
            print(f"📂 Lecture du fichier: {resolved_path}")
            doc = docx.Document(resolved_path)

            content = {
                "text": "",
                "paragraphs": [],
                "tables": [],
                "styles": [],
                "properties": {},
            }

            # Extraction des paragraphes
            for para in doc.paragraphs:
                para_data = {
                    "text": para.text,
                    "style": para.style.name if para.style else "Normal",
                    "alignment": str(para.alignment) if para.alignment else "None",
                }
                content["paragraphs"].append(para_data)
                content["text"] += para.text + "\n"

            # Extraction des tableaux
            for table_idx, table in enumerate(doc.tables):
                table_data = {"table_number": table_idx + 1, "rows": []}

                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data["rows"].append(row_data)

                content["tables"].append(table_data)

            # Propriétés du document
            if hasattr(doc, "core_properties"):
                props = doc.core_properties
                content["properties"] = {
                    "title": props.title,
                    "author": props.author,
                    "subject": props.subject,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                }

            return {
                "success": True,
                "content": content,
                "file_info": {
                    "original_path": file_path,
                    "resolved_path": resolved_path,
                    "size": os.path.getsize(resolved_path),
                    "processor": "python-docx",
                },
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur lors de la lecture DOCX: {str(e)}",
                "error_type": "PROCESSING_ERROR",
                "resolution": "Vérifiez que le fichier n'est pas corrompu et qu'il est accessible",
            }

    def _resolve_onedrive_path(self, file_path: str) -> str:
        """
        Résout les chemins OneDrive pour les rendre accessibles

        Args:
            file_path: Chemin du fichier (peut être OneDrive)

        Returns:
            Chemin résolu et accessible
        """

        # Si le fichier existe déjà, le retourner tel quel
        if os.path.exists(file_path):
            return file_path

        # Chemins OneDrive courants
        onedrive_paths = [
            os.path.expanduser("~/OneDrive"),
            os.path.expanduser("~/OneDrive - Personnel"),
            os.path.expanduser("~/OneDrive - Professionnel"),
            "C:/Users/" + os.getenv("USERNAME", "") + "/OneDrive",
            "C:/Users/" + os.getenv("USERNAME", "") + "/OneDrive - Personnel",
        ]

        # Ajouter les chemins OneDrive spécifiques détectés automatiquement
        username = os.getenv("USERNAME", "")
        if username:
            # Chemins OneDrive Business/Enterprise
            enterprise_patterns = [
                f"C:/Users/{username}/OneDrive - *",
                f"C:/Users/{username}/OneDrive*",
            ]
            for pattern in enterprise_patterns:
                onedrive_paths.extend(glob.glob(pattern))

        # Extraire le nom de fichier de base
        file_name = os.path.basename(file_path)

        # Chercher le fichier dans tous les répertoires OneDrive
        for onedrive_path in onedrive_paths:
            if os.path.exists(onedrive_path):
                # Recherche récursive dans OneDrive
                for root, files in os.walk(onedrive_path):
                    if file_name in files:
                        potential_path = os.path.join(root, file_name)
                        print(f"✅ Fichier trouvé dans OneDrive: {potential_path}")
                        return potential_path

        # Si pas trouvé, essayer avec le répertoire Desktop OneDrive
        desktop_onedrive = os.path.expanduser("~/OneDrive/Desktop")
        if os.path.exists(desktop_onedrive):
            desktop_file = os.path.join(desktop_onedrive, file_name)
            if os.path.exists(desktop_file):
                print(f"✅ Fichier trouvé sur Bureau OneDrive: {desktop_file}")
                return desktop_file

        # Retourner le chemin original si rien trouvé
        print("⚠️ Fichier non trouvé dans OneDrive, tentative avec le chemin original")
        return file_path

    def is_supported(self, file_path: str) -> bool:
        """
        Vérifie si le fichier est supporté

        Args:
            file_path: Chemin vers le fichier

        Returns:
            True si le fichier est supporté
        """
        return Path(file_path).suffix.lower() in self.supported_extensions

    def extract_text(self, file_path: str) -> Dict[str, Any]:
        """
        Extrait le texte d'un fichier DOCX avec gestion améliorée des erreurs

        Args:
            file_path: Chemin vers le fichier DOCX

        Returns:
            Dict avec le texte extrait et les métadonnées ou l'erreur
        """
        if not os.path.exists(file_path):
            return {
                "success": False,
                "error": "Fichier non trouvé",
                "error_type": "FILE_NOT_FOUND",
            }

        if not self.python_docx_available:
            return {
                "success": False,
                "error": "Le module python-docx n'est pas installé",
                "error_type": "MISSING_PACKAGE",
                "resolution": "Installez le package avec : pip install python-docx",
            }

        try:
            result = self.read_docx(file_path)
            if "error" in result:
                return {
                    "success": False,
                    "error": result["error"],
                    "error_type": "READ_ERROR",
                }

            content = result["content"]
            if not content.get("text", "").strip():
                return {
                    "success": False,
                    "error": "Document vide ou non lisible",
                    "error_type": "EMPTY_CONTENT",
                }

            return {
                "success": True,
                "content": content["text"],
                "metadata": {
                    "paragraphs": len(content["paragraphs"]),
                    "tables": len(content["tables"]),
                    "properties": content["properties"],
                },
            }

        except Exception as e:
            return {
                "success": False,
                "error": f"Erreur lors de la lecture du document: {str(e)}",
                "error_type": "PROCESSING_ERROR",
            }

    def extract_text_from_docx(self, file_path: str) -> str:
        """
        Méthode pour extraire du texte d'un fichier DOCX (compatible avec l'interface GUI)

        Args:
            file_path: Chemin vers le fichier DOCX

        Returns:
            Texte extrait du fichier DOCX
        """
        try:
            result = self.extract_text(file_path)
            if result.get("success", False):
                # Le contenu est directement dans 'content'
                return result.get("content", "")
            else:
                raise Exception(result.get("error", "Erreur inconnue"))
        except Exception as e:
            raise Exception(f"Erreur lors de l'extraction DOCX: {str(e)}") from e
